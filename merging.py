import os
from docx import Document
from pathlib import Path


def add_code_to_docx(directory, doc):
    """Recursively add code files and their content to a DOCX file."""

    for file_path in Path(directory).rglob('*'):
        if file_path.is_file() and file_path.suffix in ['.py', '.js', '.java', '.cpp',
                                                        '.txt']:  # Add more extensions if needed
            # Open and read the content of the file
            with open(file_path, 'r', encoding='utf-8') as f:
                code_content = f.read()

            # Add file name as a heading
            doc.add_heading(f"File: {file_path.name}", level=2)

            # Add the code content to the document
            doc.add_paragraph(code_content)
            doc.add_paragraph('\n' + '-' * 80 + '\n')  # Divider between code files


def generate_docx_for_repo(directory):
    """Generate a DOCX document containing all code files from the given directory."""
    doc = Document()
    doc.add_heading('Code Repository Documentation', level=1)

    # Start adding files from the selected repo/directory
    add_code_to_docx(directory, doc)

    # Save the generated document
    output_filename = f"{Path(directory).name}_code_documentation.docx"
    doc.save(output_filename)
    print(f"Documentation saved as {output_filename}")


def main():
    # User inputs the directory path (repo)
    repo_path = input("Enter the path to the repository or directory: ")
    if not os.path.isdir(repo_path):
        print("Invalid directory. Please try again.")
        return

    # Generate the documentation file
    generate_docx_for_repo(repo_path)


if __name__ == "__main__":b
    main()
